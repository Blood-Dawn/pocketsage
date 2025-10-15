"""Generate a Word document containing UML diagram PNGs."""

from __future__ import annotations

from pathlib import Path

from docx import Document  # type: ignore[import]
from docx.shared import Inches  # type: ignore[import]

BASE_DIR = Path(__file__).resolve().parents[1]
UML_DIR = BASE_DIR / "docs" / "uml"
OUTPUT_PATH = UML_DIR / "pocketsage_uml_diagrams.docx"


def _title_from(path: Path) -> str:
    name = path.stem.replace("_", " ").replace("-", " ")
    return name.title()


def build_document() -> None:
    png_paths = sorted(UML_DIR.glob("*/**/*.png"))
    if not png_paths:
        raise SystemExit("No PNG diagrams found under docs/uml")

    document = Document()
    document.add_heading("PocketSage UML Diagrams", 0)

    for path in png_paths:
        document.add_heading(_title_from(path), level=1)
        document.add_picture(str(path), width=Inches(6.5))
        document.add_page_break()

    document.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
